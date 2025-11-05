import os
import uuid
from typing import List, BinaryIO, Tuple
from app.models.document import Document, DocumentType
import PyPDF2
import docx
from PIL import Image
import librosa
from langchain.text_splitter import RecursiveCharacterTextSplitter


class IngestionService:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )

    async def process_file(self, file: BinaryIO, filename: str, user_id: str) -> Tuple[Document, List[str]]:
        """Process uploaded file and extract content"""

        file_ext = filename.split('.')[-1].lower()
        doc_id = str(uuid.uuid4())

        # Extract text based on file type
        if file_ext == 'pdf':
            content = self._extract_pdf(file)
            doc_type = DocumentType.PDF
        elif file_ext in ['txt', 'md']:
            content = file.read().decode('utf-8')
            doc_type = DocumentType.TEXT
        elif file_ext == 'docx':
            content = self._extract_docx(file)
            doc_type = DocumentType.DOCX
        elif file_ext in ['py', 'js', 'java', 'cpp', 'c']:
            content = file.read().decode('utf-8')
            doc_type = DocumentType.CODE
        elif file_ext in ['jpg', 'png', 'jpeg']:
            content = await self._extract_image(file)
            doc_type = DocumentType.IMAGE
        elif file_ext in ['mp3', 'wav']:
            content = await self._extract_audio(file)
            doc_type = DocumentType.AUDIO
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")

        # Split into chunks
        chunks = self.text_splitter.split_text(content)

        document = Document(
            id=doc_id,
            user_id=user_id,
            filename=filename,
            file_type=doc_type,
            size=len(content),
            chunks=[str(uuid.uuid4()) for _ in chunks],
            metadata={
                'extension': file_ext,
                'num_chunks': len(chunks)
            }
        )

        return document, chunks

    def _extract_pdf(self, file: BinaryIO) -> str:
        """Extract text from PDF"""
        pdf_reader = PyPDF2.PdfReader(file)
        text = []
        for page in pdf_reader.pages:
            text.append(page.extract_text() or "")
        return '\n'.join(text)

    def _extract_docx(self, file: BinaryIO) -> str:
        """Extract text from DOCX"""
        doc = docx.Document(file)
        return '\n'.join([para.text for para in doc.paragraphs])

    async def _extract_image(self, file: BinaryIO) -> str:
        """Extract text/description from image using vision model (placeholder)"""
        image = Image.open(file)
        return f"Image description: {image.size}"

    async def _extract_audio(self, file: BinaryIO) -> str:
        """Transcribe audio using Whisper or similar (placeholder)"""
        temp_path = f"/tmp/{uuid.uuid4()}.wav"
        with open(temp_path, 'wb') as f:
            f.write(file.read())

        y, sr = librosa.load(temp_path, sr=None)
        os.remove(temp_path)
        return f"Audio transcription placeholder (duration: {len(y)/sr}s)"
